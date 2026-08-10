from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


AFFILIATION = (
    "Department of Orthopedics Center, The First Affiliated Hospital of "
    "Shihezi University, Shihezi University, Shihezi 832008, China"
)


def set_font(document: Document) -> None:
    for style_name in ("Normal", "Title"):
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11 if style_name == "Normal" else 16)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(11)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
                        run.font.size = Pt(10)


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def build(source: Path, output: Path) -> None:
    document = Document(source)
    values = {
        "Corresponding author": "Corresponding author: Lei Chen",
        "Affiliation": f"Affiliation: {AFFILIATION}",
        "Email": "Email: 564386249@qq.com",
        "Telephone": "Telephone: 13579758836",
        "Fax": "Fax: Not provided",
    }
    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        for prefix, replacement in values.items():
            if stripped.lower().startswith(prefix.lower()):
                replace_paragraph(paragraph, replacement)
                break

    if not document.tables:
        raise RuntimeError("Author-information template does not contain a table")
    table = document.tables[0]
    if len(table.rows) < 6 or len(table.columns) < 5:
        raise RuntimeError("Unexpected author-information table structure")
    rows = [
        ("1 (co-first; equal contribution)", "Junhui", "Shi", AFFILIATION, "2223727941@qq.com"),
        ("2 (co-first; equal contribution)", "Mengxiang", "Liu", AFFILIATION, "2877992646@qq.com"),
        ("3", "Repkat", "Inayatilla", AFFILIATION, "2047733903@qq.com"),
        ("4", "Ke", "Li", AFFILIATION, "1413458714@qq.com"),
        ("5 (corresponding author)", "Lei", "Chen", AFFILIATION, "564386249@qq.com"),
    ]
    for row, values_row in zip(table.rows[1:6], rows):
        for cell, value in zip(row.cells[:5], values_row):
            cell.text = value

    summary = document.add_paragraph()
    summary.add_run("Author sequence: ").bold = True
    summary.add_run("Junhui Shi, Mengxiang Liu, Repkat Inayatilla, Ke Li, Lei Chen. ")
    summary.add_run("Equal-contribution statement: ").bold = True
    summary.add_run("Junhui Shi and Mengxiang Liu contributed equally and share first authorship.")
    summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(0)
    summary.paragraph_format.line_spacing = 1.0

    set_font(document)
    for run in summary.runs:
        run.font.size = Pt(9)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Wrote {output}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build(args.source, args.output)


if __name__ == "__main__":
    main()
