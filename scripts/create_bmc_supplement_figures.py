from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FIGURES = [
    (
        "Figure S1. Differential-expression threshold sensitivity",
        "SupplementaryFigure1_DEG_threshold_sensitivity.png",
        "Shared-gene retention across six prespecified combinations of FDR and absolute log2-fold-change thresholds. The highlighted point denotes the primary threshold. Complete membership is provided in Additional file 1: Tables S3a and S3b.",
    ),
    (
        "Figure S2. Hallmark states in discovery and external tissue cohorts",
        "SupplementaryFigure2_external_tissue_Hallmark.png",
        "Normalized enrichment scores for the ten Hallmark pathways significant in both discovery cohorts. Crosses mark FDR ≥0.05 or an unavailable enrichment estimate. The matrix shows strong OC replication and greater OA cohort dependence.",
    ),
    (
        "Figure S3. Dataset-specific single-cell embeddings",
        "SupplementaryFigure3_all_single_cell_UMAPs.png",
        "All five single-cell atlases with exact source labels. OA and OC datasets were retained as separate analytical spaces.",
    ),
    (
        "Figure S4. Discovery-cohort PCA and sample-correlation quality control",
        "SupplementaryFigure4_bulk_PCA_QC.png",
        "Separate unsupervised audits for the OA and OC discovery cohorts. These displays assessed cohort structure but were not used for outcome-informed sample removal or batch correction.",
    ),
    (
        "Figure S5. Cross-layer evidence summary for the illustrative genes",
        "SupplementaryFigure5_illustrative_gene_evidence.png",
        "Descriptive summary of external-direction agreement, source-defined single-cell localization, dual-blood FDR status and interpretive role for G0S2, EFEMP1, AKAP12, SOX9 and DDIT3. These genes illustrate distinct evidence patterns and do not constitute an optimized predictive signature.",
    ),
]


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    for style_name in ("Normal", "Title", "Heading 1"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Title"].font.size = Pt(17)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(12)
    doc.styles["Heading 1"].font.bold = True


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--figure-dir", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    doc = Document()
    configure(doc)
    title = doc.add_paragraph(style="Title")
    title.add_run("Additional file 2: Supplementary figures")
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Shared molecular features between osteoarthritis and ovarian cancer revealed by multi-layer transcriptomic analyses"
    ).italic = True
    doc.add_paragraph(
        "This file contains five supplementary figures cited in the main manuscript."
    )

    for index, (heading, filename, legend) in enumerate(FIGURES):
        if index > 0:
            doc.add_page_break()
        paragraph = doc.add_paragraph(heading, style="Heading 1")
        paragraph.paragraph_format.keep_with_next = True
        image_path = args.figure_dir / filename
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        picture_paragraph = doc.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(image_path), width=Cm(17.0))
        caption = doc.add_paragraph(legend)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.keep_together = True

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
