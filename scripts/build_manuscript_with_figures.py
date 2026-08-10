from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FIGURES = [
    ("Figure 1", "Figure1_study_design.png"),
    ("Figure 2", "Figure2_bulk_discovery.png"),
    ("Figure 3", "Figure3_network_and_ml_stability.png"),
    ("Figure 4", "Figure4_external_validation.png"),
    ("Figure 5", "Figure5_single_cell_localization.png"),
    ("Figure 6", "Figure6_functional_and_tcga_context.png"),
    ("Supplementary Figure 1", "SupplementaryFigure1_sensitivity_details.png"),
    ("Supplementary Figure 2", "SupplementaryFigure2_negative_bidirectional_MR.png"),
    ("Supplementary Figure 3", "SupplementaryFigure3_single_cell_UMAPs.png"),
    (
        "Supplementary Figure 4",
        "SupplementaryFigure4_HPA_normal_tissue_context.png",
    ),
    (
        "Supplementary Figure 5",
        "SupplementaryFigure5_TCGA_relative_context.png",
    ),
]


def parse_legends(path: Path) -> dict[str, tuple[str, str]]:
    text = path.read_text(encoding="utf-8")
    matches = list(re.finditer(r"^## (.+)$", text, flags=re.MULTILINE))
    legends: dict[str, tuple[str, str]] = {}
    for index, match in enumerate(matches):
        heading = match.group(1).strip()
        body_start = match.end()
        body_end = (
            matches[index + 1].start()
            if index + 1 < len(matches)
            else len(text)
        )
        body = text[body_start:body_end].strip()
        body = re.sub(r"\*\*(.*?)\*\*", r"\1", body)
        body = re.sub(r"\s+", " ", body).strip()
        label, _ = heading.split(". ", 1)
        legends[label] = (heading, body)
    return legends


def set_font(run, name: str, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def build(
    source_docx: Path,
    figure_dir: Path,
    legends_md: Path,
    output_docx: Path,
) -> None:
    legends = parse_legends(legends_md)
    document = Document(source_docx)
    section = document.sections[-1]
    available_width = (
        section.page_width - section.left_margin - section.right_margin
    )
    figure_width = min(available_width, Inches(6.35))

    document.add_page_break()
    section_heading = document.add_paragraph(style="Heading 1")
    section_heading.paragraph_format.keep_with_next = True
    section_heading.add_run("Embedded figures")

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    intro.add_run(
        "Main and supplementary figures are embedded below at "
        "publication-review resolution. The preceding figure-legends section "
        "remains the authoritative caption set."
    )

    for index, (label, filename) in enumerate(FIGURES):
        image_path = figure_dir / filename
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        if label not in legends:
            raise KeyError(f"Missing legend for {label}")
        if index > 0:
            document.add_page_break()

        heading, caption = legends[label]
        title_paragraph = document.add_paragraph(style="Heading 2")
        title_paragraph.paragraph_format.keep_with_next = True
        title_paragraph.paragraph_format.space_after = Pt(7)
        title_paragraph.add_run(heading)

        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.keep_with_next = True
        image_paragraph.paragraph_format.space_after = Pt(6)
        picture_run = image_paragraph.add_run()
        picture_run.add_picture(str(image_path), width=figure_width)
        inline = picture_run._r.xpath(".//wp:inline")[0]
        doc_pr = inline.xpath("./wp:docPr")[0]
        doc_pr.set("title", heading)
        doc_pr.set("descr", caption)

        caption_paragraph = document.add_paragraph(style="Caption")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption_paragraph.paragraph_format.space_before = Pt(0)
        caption_paragraph.paragraph_format.space_after = Pt(0)
        caption_paragraph.paragraph_format.line_spacing = 1.05
        caption_run = caption_paragraph.add_run(caption)
        set_font(caption_run, "Arial", 9.0, color="333333")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)
    print(f"Wrote {output_docx}")
    print(f"Embedded {len(FIGURES)} figures")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--figure-dir", required=True, type=Path)
    parser.add_argument("--legends", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build(
        args.source.resolve(),
        args.figure_dir.resolve(),
        args.legends.resolve(),
        args.output.resolve(),
    )


if __name__ == "__main__":
    main()
