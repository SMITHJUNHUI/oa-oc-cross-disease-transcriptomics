#!/usr/bin/env python3
"""Create a restrained journal-style DOCX from the versioned manuscript Markdown."""

from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ACCENT = "176B73"
ACCENT_DARK = "164E57"
INK = "1D2935"
MUTED = "5B6770"
PALE = "EAF4F4"
REVIEW = "FFF3CD"
WHITE = "FFFFFF"
FONT_SANS = "Arial"
FONT_SERIF = "Times New Roman"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_rows_disallowed(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def set_run_font(run, name: str, size: float | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text: str) -> None:
    """Add a small, deterministic Markdown inline subset."""
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Consolas", 9.0, ACCENT_DARK)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=False, page_break_before=False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (
        ("keepNext", keep_next),
        ("keepLines", keep_lines),
        ("pageBreakBefore", page_break_before),
    ):
        if enabled:
            p_pr.append(OxmlElement(f"w:{tag}"))


def add_bottom_border(paragraph, color=ACCENT, size=12, space=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_document_language(doc: Document, language: str = "en-US") -> None:
    styles = doc.styles
    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        r_pr = style.element.get_or_add_rPr()
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:val"), language)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_SERIF
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SERIF)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = FONT_SANS
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.line_spacing = 1.05

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT_SANS
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
    subtitle.font.size = Pt(10)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)

    for style_name, size, before, after in (
        ("Heading 1", 15, 18, 7),
        ("Heading 2", 12, 12, 4),
        ("Heading 3", 10.5, 9, 3),
    ):
        style = styles[style_name]
        style.font.name = FONT_SANS
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Manuscript Metadata" not in styles:
        metadata = styles.add_style("Manuscript Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata = styles["Manuscript Metadata"]
    metadata.font.name = FONT_SANS
    metadata._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SANS)
    metadata.font.size = Pt(10)
    metadata.font.color.rgb = RGBColor.from_string(INK)
    metadata.paragraph_format.space_after = Pt(4)
    metadata.paragraph_format.line_spacing = 1.05

    if "Reference" not in styles:
        reference = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = styles["Reference"]
    reference.font.name = FONT_SERIF
    reference._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SERIF)
    reference.font.size = Pt(9)
    reference.font.color.rgb = RGBColor.from_string(INK)
    reference.paragraph_format.left_indent = Cm(0.7)
    reference.paragraph_format.first_line_indent = Cm(-0.7)
    reference.paragraph_format.space_after = Pt(3)
    reference.paragraph_format.line_spacing = 1.05

    if "Figure Legend" not in styles:
        legend = styles.add_style("Figure Legend", WD_STYLE_TYPE.PARAGRAPH)
    else:
        legend = styles["Figure Legend"]
    legend.font.name = FONT_SERIF
    legend._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SERIF)
    legend.font.size = Pt(9.5)
    legend.font.color.rgb = RGBColor.from_string(INK)
    legend.paragraph_format.space_after = Pt(8)
    legend.paragraph_format.line_spacing = 1.15

    set_document_language(doc)


def configure_section(section, first=False) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)
    if first:
        section.different_first_page_header_footer = True


def add_header_footer(section, revision_label: str = "V2") -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(
        f"OA–OC TRANSCRIPTOMIC CONTEXTS  •  REVISION {revision_label.upper()}"
    )
    set_run_font(run, FONT_SANS, 8.0, MUTED)
    run.bold = True
    add_bottom_border(paragraph, color="B9D9DB", size=6, space=3)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, FONT_SANS, 8.0, MUTED)
    add_page_number(paragraph)


def add_status_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, REVIEW)
    set_cell_margins(cell, 140, 180, 140, 180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("AUTHOR REVIEW REQUIRED")
    set_run_font(run, FONT_SANS, 9.0, "7A5B00")
    run.bold = True
    paragraph.add_run(
        "\nComplete author, affiliation, journal, funding, competing-interest, "
        "ethics-policy, and contribution fields before submission."
    )
    for run in paragraph.runs[1:]:
        set_run_font(run, FONT_SANS, 9.0, INK)


def add_title_page(
    doc: Document,
    title: str,
    metadata_lines: list[str],
    revision_label: str = "V2",
    subtitle_text: str | None = None,
) -> None:
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = kicker.add_run("ORIGINAL RESEARCH  /  GENERAL JOURNAL FORMAT")
    set_run_font(run, FONT_SANS, 9.5, ACCENT)
    run.bold = True
    kicker.paragraph_format.space_after = Pt(18)

    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(title)
    add_bottom_border(paragraph, color=ACCENT, size=18, space=10)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(subtitle_text or (
        "Cross-disease analysis with external tissue replication, cellular "
        "localization, and conditional blood validation"
    ))
    subtitle.paragraph_format.space_after = Pt(18)

    for line in metadata_lines:
        paragraph = doc.add_paragraph(style="Manuscript Metadata")
        add_inline(paragraph, line)

    doc.add_paragraph()
    add_status_box(doc)

    stamp = doc.add_paragraph()
    stamp.paragraph_format.space_before = Pt(12)
    stamp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = stamp.add_run(
        f"Prepared {date.today().strftime('%d %B %Y')}  •  Version {revision_label}"
    )
    set_run_font(run, FONT_SANS, 8.5, MUTED)

    doc.add_page_break()


def is_front_matter(line: str) -> bool:
    return line.startswith("**") and ":**" in line


def parse_markdown(markdown_path: Path) -> tuple[str, list[str], list[str]]:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("Manuscript Markdown must start with a level-1 title")
    title = lines[0][2:].strip()
    metadata: list[str] = []
    body_start = 1
    for idx in range(1, len(lines)):
        line = lines[idx].strip()
        if not line:
            continue
        if is_front_matter(line):
            metadata.append(line)
            body_start = idx + 1
            continue
        if line.startswith("## "):
            body_start = idx
            break
    return title, metadata, lines[body_start:]


def add_body(doc: Document, lines: list[str]) -> None:
    in_references = False
    in_legends = False
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            in_references = heading == "References"
            in_legends = heading == "Figure legends"
            paragraph = doc.add_paragraph(heading, style="Heading 1")
            if heading in {"References", "Figure legends", "Supplementary table index"}:
                set_paragraph_keep(paragraph, keep_next=True, page_break_before=True)
            else:
                set_paragraph_keep(paragraph, keep_next=True)
            add_bottom_border(paragraph, color="B9D9DB", size=8, space=4)
            idx += 1
            continue

        if line.startswith("### "):
            heading = line[4:].strip()
            paragraph = doc.add_paragraph(heading, style="Heading 2")
            set_paragraph_keep(paragraph, keep_next=True)
            idx += 1
            continue

        if line.startswith("#### "):
            paragraph = doc.add_paragraph(line[5:].strip(), style="Heading 3")
            set_paragraph_keep(paragraph, keep_next=True)
            idx += 1
            continue

        if in_references and re.match(r"^\d+\.\s", line):
            paragraph = doc.add_paragraph(style="Reference")
            add_inline(paragraph, line)
            set_paragraph_keep(paragraph, keep_lines=True)
            idx += 1
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Cm(0.65)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline(paragraph, line[2:])
            set_paragraph_keep(paragraph, keep_lines=True)
            idx += 1
            continue

        paragraph_style = "Figure Legend" if in_legends else "Normal"
        paragraph = doc.add_paragraph(style=paragraph_style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(paragraph, line)
        set_paragraph_keep(paragraph, keep_lines=True)
        idx += 1


def build_document(
    markdown_path: Path,
    output_path: Path,
    revision_label: str = "V2",
    subtitle_text: str | None = None,
) -> None:
    title, metadata, body = parse_markdown(markdown_path)
    doc = Document()
    configure_styles(doc)
    for index, section in enumerate(doc.sections):
        configure_section(section, first=index == 0)
        add_header_footer(section, revision_label)

    doc.core_properties.title = title
    doc.core_properties.subject = (
        "Reproducible osteoarthritis–ovarian cancer transcriptomic study; "
        f"revision {revision_label}"
    )
    doc.core_properties.author = "Author details pending"
    doc.core_properties.keywords = (
        "osteoarthritis, ovarian cancer, transcriptomics, single-cell RNA sequencing, "
        "WGCNA, peripheral blood, reproducibility"
    )
    doc.core_properties.comments = (
        "Generated from the versioned Markdown source. Author and journal metadata require review."
    )

    add_title_page(doc, title, metadata, revision_label, subtitle_text)
    add_body(doc, body)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--revision-label", default="V2")
    parser.add_argument("--subtitle")
    args = parser.parse_args()
    build_document(
        args.input.resolve(),
        args.output.resolve(),
        args.revision_label,
        args.subtitle,
    )
    print(f"Wrote {args.output.resolve()}")


if __name__ == "__main__":
    main()
