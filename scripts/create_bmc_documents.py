from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT = "Arial"
BLACK = "000000"
PLACEHOLDER_FILL = "FFF2CC"


def set_run_font(run, size: float = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(BLACK)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def highlight_placeholder(run) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PLACEHOLDER_FILL)
    run._element.get_or_add_rPr().append(shading)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("Page ")
    set_run_font(run, 10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    element = OxmlElement("w:lnNumType")
    element.set(qn("w:countBy"), "1")
    element.set(qn("w:distance"), "360")
    element.set(qn("w:restart"), "continuous")
    sect_pr.append(element)


def configure_document(doc: Document, line_numbers: bool = True, page_numbers: bool = True) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    if line_numbers:
        configure_line_numbers(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(BLACK)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    title.paragraph_format.space_after = Pt(6)
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True

    for style_name in ("Reference", "Figure Legend", "Metadata"):
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.space_after = Pt(0)

    reference = styles["Reference"]
    reference.paragraph_format.left_indent = Cm(0.8)
    reference.paragraph_format.first_line_indent = Cm(-0.8)

    if page_numbers:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer)


def add_inline(paragraph, text: str, size: float = 12) -> None:
    token_pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|\[AUTHOR INPUT REQUIRED:.*?\]|\[[A-Z][A-Z _,-]*\])")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size, bold=True)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size, italic=True)
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size, bold=True)
            highlight_placeholder(run)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size)


def parse_manuscript(markdown: str) -> tuple[str, list[str], list[str]]:
    lines = markdown.splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("A level-one manuscript title is required")
    title = lines[0][2:].strip()
    metadata: list[str] = []
    body_start = 1
    for index in range(1, len(lines)):
        line = lines[index].strip()
        if not line:
            continue
        if line.startswith("**") and ":**" in line:
            metadata.append(line)
            body_start = index + 1
            continue
        if line.startswith("## "):
            body_start = index
            break
    return title, metadata, lines[body_start:]


def build_manuscript(input_path: Path, output_path: Path) -> None:
    title, metadata, lines = parse_manuscript(input_path.read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc, line_numbers=True)
    doc.core_properties.title = title
    doc.core_properties.subject = "Research Article prepared for BMC Medical Genomics"
    doc.core_properties.keywords = "osteoarthritis, ovarian cancer, transcriptomics, external validation, single-cell RNA sequencing, peripheral blood"

    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(title_paragraph, title)
    title_ppr = title_paragraph._p.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    for line in metadata:
        paragraph = doc.add_paragraph(style="Metadata")
        add_inline(paragraph, line)

    in_references = False
    in_legends = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if in_legends and re.match(r"Figure \d+\.", heading):
                paragraph = doc.add_paragraph(heading, style="Heading 2")
            else:
                paragraph = doc.add_paragraph(heading, style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            in_references = heading == "References"
            in_legends = heading == "Figure legends" or (in_legends and heading.startswith("Figure "))
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(line[4:].strip(), style="Heading 2")
            paragraph.paragraph_format.keep_with_next = True
            continue
        if line.startswith("#### "):
            paragraph = doc.add_paragraph(line[5:].strip(), style="Heading 3")
            paragraph.paragraph_format.keep_with_next = True
            continue
        if in_references and re.match(r"^\d+\.\s", line):
            paragraph = doc.add_paragraph(style="Reference")
        elif in_legends:
            paragraph = doc.add_paragraph(style="Figure Legend")
        else:
            paragraph = doc.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(paragraph, line)
        paragraph.paragraph_format.keep_together = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_cover_letter(input_path: Path, output_path: Path) -> None:
    lines = input_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc, line_numbers=False, page_numbers=False)
    doc.core_properties.title = "Cover letter to BMC Medical Genomics"
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            continue
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(4)
        add_inline(paragraph, stripped.replace("  ", ""), size=11)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manuscript-md", required=True, type=Path)
    parser.add_argument("--manuscript-docx", required=True, type=Path)
    parser.add_argument("--cover-md", required=True, type=Path)
    parser.add_argument("--cover-docx", required=True, type=Path)
    args = parser.parse_args()
    build_manuscript(args.manuscript_md, args.manuscript_docx)
    build_cover_letter(args.cover_md, args.cover_docx)
    print(f"Wrote {args.manuscript_docx}")
    print(f"Wrote {args.cover_docx}")


if __name__ == "__main__":
    main()
